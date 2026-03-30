#!/usr/bin/env python3
"""
Legal document -> style/voice dataset generator.

Builds two JSONL outputs from one file or a directory of legal documents:
1. style-analysis records with paragraph-level metadata and style features
2. training pairs that abstract case-specific facts while preserving legal voice

Supported inputs:
- PDF
- DOCX
- TXT / MD

OCR support:
- PDF OCR can be enabled with --ocr-mode auto|always|never
- OCR requires the Python packages in requirements-legal.txt and the tesseract
  binary to be installed on the host machine

Example:
  python3 legal_style_dataset.py ./US_v_Trump_23_cr_257.pdf ./out/jack_smith \
      --author "Jack Smith" --document-type indictment --ocr-mode auto
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import fitz
import pytesseract
from docx import Document as DocxDocument
from PIL import Image
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
MONTH_PATTERN = (
    r"January|February|March|April|May|June|July|August|September|October|"
    r"November|December|Jan\.|Feb\.|Mar\.|Apr\.|Jun\.|Jul\.|Aug\.|Sep\.|"
    r"Sept\.|Oct\.|Nov\.|Dec\."
)
TRANSITION_PHRASES = [
    "further", "furthermore", "moreover", "thereafter", "therefore",
    "accordingly", "however", "instead", "because", "as a result",
    "for example", "in particular", "in turn", "then", "next",
    "at all times material", "from on or about", "in fact", "thus",
]
QUALIFIER_WORDS = [
    "knowingly", "willfully", "approximately", "material", "false",
    "alleged", "unlawful", "lawful", "corruptly", "intentionally",
    "thereafter", "therein", "herein", "respectively",
]
PASSIVE_MARKERS = {"was", "were", "is", "are", "be", "been", "being"}
LEGAL_ENTITY_EXCLUSIONS = {
    "The Defendant", "The Government", "The Grand Jury", "The Court",
    "District of Columbia", "United States", "United States of America",
    "Count One", "Count Two", "Count Three", "Count Four", "Count Five",
    "Count Six", "Count Seven", "Count Eight", "Count Nine", "Count Ten",
}


@dataclass
class ParagraphUnit:
    source_file: str
    paragraph_id: str
    author: str
    document_type: str
    page_start: int | None
    ordinal: int
    section_path: list[str]
    text: str


def normalize_whitespace(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = text.replace("\u2019", "'")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_line(line: str) -> str:
    line = normalize_whitespace(line)
    line = re.sub(r"\s+([,.;:])", r"\1", line)
    return line.strip()


def normalize_paragraph_text(text: str) -> str:
    text = normalize_whitespace(text)
    text = re.sub(r"(\w)-\s+(\w)", r"\1\2", text)
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    text = re.sub(r"\(\s+", "(", text)
    text = re.sub(r"\s+\)", ")", text)
    return text.strip()


def slugify(value: str) -> str:
    value = value.lower()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-") or "document"


def detect_document_type(path: Path, override: str | None, preview_text: str) -> str:
    if override:
        return override

    text = f"{path.stem} {preview_text[:4000]}".lower()
    checks = [
        ("indictment", "indictment"),
        ("complaint", "complaint"),
        ("motion", "motion"),
        ("brief", "brief"),
        ("memorandum", "memorandum"),
        ("opinion", "opinion"),
        ("order", "order"),
        ("petition", "petition"),
    ]
    for needle, label in checks:
        if needle in text:
            return label
    return "legal_document"


def collect_input_files(input_path: Path) -> list[Path]:
    if input_path.is_file():
        if input_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {input_path.suffix}")
        return [input_path]

    files = [
        path for path in sorted(input_path.rglob("*"))
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    if not files:
        raise ValueError(f"No supported files found in {input_path}")
    return files


def repeated_header_or_footer(line: str) -> bool:
    lowered = line.lower()
    if re.match(r"^case\s+.+document\s+\d+\s+filed\s+.+page\s+\d+\s+of\s+\d+", lowered):
        return True
    if re.match(r"^page\s+\d+\s+of\s+\d+$", lowered):
        return True
    if re.match(r"^\d+$", lowered):
        return True
    return False


def is_heading(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if len(stripped.split()) > 12:
        return False
    if stripped.endswith((",", ".")):
        return False
    if re.match(r"^COUNT\s+[A-Z0-9]+$", stripped):
        return True
    if re.match(r"^\([A-Za-z0-9]+\)$", stripped):
        return False
    if stripped.isupper() and any(char.isalpha() for char in stripped):
        return True
    if re.match(r"^[A-Z][A-Za-z ]+$", stripped) and not stripped.endswith((".", ";", ":")):
        return True
    return False


def update_section_path(section_path: list[str], heading: str) -> list[str]:
    if re.match(r"^COUNT\s+[A-Z0-9]+$", heading):
        return [heading]
    if heading.isupper():
        if section_path and section_path[-1] == heading:
            return section_path
        return (section_path[:1] + [heading])[-3:]
    return (section_path + [heading])[-3:]


def split_pdf_page_into_chunks(page_text: str) -> list[str]:
    chunks: list[str] = []
    current: list[str] = []

    for raw_line in page_text.splitlines():
        line = normalize_line(raw_line)
        if not line or repeated_header_or_footer(line):
            if current:
                chunks.append(" ".join(current).strip())
                current = []
            continue

        if is_heading(line):
            if current:
                chunks.append(" ".join(current).strip())
                current = []
            chunks.append(line)
            continue

        if current and re.match(r"^\d+[.)]\s+", line):
            chunks.append(" ".join(current).strip())
            current = [line]
            continue

        current.append(line)

    if current:
        chunks.append(" ".join(current).strip())

    return [chunk for chunk in chunks if chunk.strip()]


def starts_new_paragraph(text: str) -> bool:
    return bool(
        re.match(r"^\d+[.)]\s+", text)
        or re.match(r"^[A-Z][A-Z ]+$", text)
        or re.match(r"^(The\s+Conspiracy|Manner and Means|Purpose of the Conspiracy)\b", text)
    )


def ends_paragraph(text: str) -> bool:
    return text.rstrip().endswith((".", "!", "?", ":"))


def split_text_into_chunks(text: str) -> list[str]:
    raw_chunks = re.split(r"\n\s*\n+", normalize_whitespace(text))
    chunks: list[str] = []
    for raw_chunk in raw_chunks:
        lines = [normalize_line(line) for line in raw_chunk.splitlines() if normalize_line(line)]
        if not lines:
            continue
        if len(lines) == 1:
            chunks.append(lines[0])
            continue
        chunks.append(" ".join(lines))
    return chunks


def ocr_pdf_page(pdf_doc: fitz.Document, page_index: int, ocr_lang: str) -> str:
    page = pdf_doc.load_page(page_index)
    pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
    image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
    return normalize_paragraph_text(pytesseract.image_to_string(image, lang=ocr_lang))


def extract_pdf_page_texts(path: Path, ocr_mode: str, ocr_min_chars: int, ocr_lang: str) -> tuple[list[tuple[int, str]], bool, list[int]]:
    reader = PdfReader(str(path))
    fitz_doc = fitz.open(path)
    tesseract_path = shutil.which("tesseract")

    page_texts: list[tuple[int, str]] = []
    ocr_used = False
    ocr_pages: list[int] = []

    for page_index, page in enumerate(reader.pages):
        extracted = normalize_whitespace(page.extract_text() or "")
        should_ocr = ocr_mode == "always" or (ocr_mode == "auto" and len(extracted) < ocr_min_chars)

        if should_ocr and not tesseract_path:
            if ocr_mode == "always":
                raise RuntimeError("OCR mode is 'always' but tesseract is not installed")
            should_ocr = False

        if should_ocr:
            ocr_text = ocr_pdf_page(fitz_doc, page_index, ocr_lang)
            if len(ocr_text) > len(extracted):
                extracted = ocr_text
                ocr_used = True
                ocr_pages.append(page_index + 1)

        page_texts.append((page_index + 1, extracted))

    fitz_doc.close()
    return page_texts, ocr_used, ocr_pages


def make_document_id(path: Path) -> str:
    return hashlib.sha1(str(path).encode("utf-8")).hexdigest()[:12]


def parse_pdf_document(path: Path, author: str, document_type: str, ocr_mode: str,
                       ocr_min_chars: int, ocr_lang: str) -> tuple[list[ParagraphUnit], dict]:
    page_texts, ocr_used, ocr_pages = extract_pdf_page_texts(path, ocr_mode, ocr_min_chars, ocr_lang)
    section_path: list[str] = []
    units: list[ParagraphUnit] = []
    doc_id = make_document_id(path)
    ordinal = 0
    carry_text = ""
    carry_page_start: int | None = None

    def flush_carry() -> None:
        nonlocal carry_text, carry_page_start, ordinal
        text = normalize_paragraph_text(carry_text)
        if not text:
            carry_text = ""
            carry_page_start = None
            return
        ordinal += 1
        units.append(
            ParagraphUnit(
                source_file=str(path.name),
                paragraph_id=f"{doc_id}-p{ordinal:04d}",
                author=author,
                document_type=document_type,
                page_start=carry_page_start,
                ordinal=ordinal,
                section_path=list(section_path),
                text=text,
            )
        )
        carry_text = ""
        carry_page_start = None

    for page_number, page_text in page_texts:
        for chunk in split_pdf_page_into_chunks(page_text):
            text = normalize_paragraph_text(chunk)
            if not text:
                continue
            if is_heading(text):
                flush_carry()
                section_path = update_section_path(section_path, text)
                continue

            if not carry_text:
                carry_text = text
                carry_page_start = page_number
                continue

            if not ends_paragraph(carry_text) and not starts_new_paragraph(text):
                carry_text = f"{carry_text} {text}".strip()
                continue

            flush_carry()
            carry_text = text
            carry_page_start = page_number

    flush_carry()

    meta = {
        "document_id": doc_id,
        "ocr_used": ocr_used,
        "ocr_pages": ocr_pages,
        "pages": len(page_texts),
    }
    return units, meta


def parse_docx_document(path: Path, author: str, document_type: str) -> tuple[list[ParagraphUnit], dict]:
    document = DocxDocument(str(path))
    units: list[ParagraphUnit] = []
    section_path: list[str] = []
    doc_id = make_document_id(path)
    ordinal = 0

    for paragraph in document.paragraphs:
        text = normalize_paragraph_text(paragraph.text)
        if not text:
            continue
        if is_heading(text):
            section_path = update_section_path(section_path, text)
            continue
        ordinal += 1
        units.append(
            ParagraphUnit(
                source_file=str(path.name),
                paragraph_id=f"{doc_id}-p{ordinal:04d}",
                author=author,
                document_type=document_type,
                page_start=None,
                ordinal=ordinal,
                section_path=list(section_path),
                text=text,
            )
        )

    return units, {"document_id": doc_id, "ocr_used": False, "ocr_pages": [], "pages": None}


def parse_text_document(path: Path, author: str, document_type: str) -> tuple[list[ParagraphUnit], dict]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    units: list[ParagraphUnit] = []
    section_path: list[str] = []
    doc_id = make_document_id(path)
    ordinal = 0

    for chunk in split_text_into_chunks(text):
        normalized = normalize_paragraph_text(chunk)
        if not normalized:
            continue
        if is_heading(normalized):
            section_path = update_section_path(section_path, normalized)
            continue
        ordinal += 1
        units.append(
            ParagraphUnit(
                source_file=str(path.name),
                paragraph_id=f"{doc_id}-p{ordinal:04d}",
                author=author,
                document_type=document_type,
                page_start=None,
                ordinal=ordinal,
                section_path=list(section_path),
                text=normalized,
            )
        )

    return units, {"document_id": doc_id, "ocr_used": False, "ocr_pages": [], "pages": None}


def load_preview_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")[:4000]
    if suffix == ".docx":
        document = DocxDocument(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs[:20])[:4000]
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join((reader.pages[index].extract_text() or "") for index in range(min(3, len(reader.pages))))[:4000]
    return ""


def split_sentences(text: str) -> list[str]:
    protected = re.sub(r"(^|\s)(\d+)\.\s+(?=[A-Z\[])", r"\1\2<ENUM> ", text)
    parts = re.split(r"(?<=[.!?])\s+(?=[A-Z\[])", protected)
    parts = [part.replace("<ENUM>", ".") for part in parts]
    return [part.strip() for part in parts if part.strip()]


def extract_transition_phrases(text: str) -> list[str]:
    lowered = text.lower()
    return [phrase for phrase in TRANSITION_PHRASES if phrase in lowered]


def classify_rhetorical_move(unit: ParagraphUnit, text: str) -> str:
    lowered = text.lower()
    section_text = " ".join(unit.section_path).lower()

    if any(token in lowered for token in ["united states district court", "defendant.", "grand jury original"]):
        return "caption"
    if "allegations contained in paragraphs" in lowered or "re-alleged and fully incorporated" in lowered:
        return "incorporation_by_reference"
    if "violation" in lowered and "u.s.c." in lowered:
        return "count_charge"
    if "from on or about" in lowered or len(re.findall(MONTH_PATTERN, text, flags=re.IGNORECASE)) >= 2:
        return "chronology"
    if "the purpose of the conspiracy" in lowered or "the object of the conspiracy" in lowered:
        return "theory_of_liability"
    if "therefore" in lowered or lowered.startswith("thus"):
        return "conclusion"
    if "court" in section_text or "standard" in section_text:
        return "legal_standard"
    if "count" in section_text:
        return "factual_allegation"
    if text.startswith(tuple(str(number) + "." for number in range(1, 100))):
        return "numbered_allegation"
    return "factual_allegation"


def compute_style_features(text: str, rhetorical_move: str) -> dict:
    sentences = split_sentences(text)
    words = re.findall(r"[A-Za-z][A-Za-z'-]*", text)
    avg_sentence_length = round(len(words) / max(len(sentences), 1), 2)
    citation_count = len(re.findall(r"\b\d+\s+U\.S\.C\.\s+§{1,2}\s*[\w().-]+", text))
    citation_count += len(re.findall(r"\b\d+\s+(?:F\.|S\. Ct\.|U\.S\.)", text))
    date_count = len(re.findall(MONTH_PATTERN, text, flags=re.IGNORECASE))
    qualifiers = [word for word in QUALIFIER_WORDS if re.search(rf"\b{re.escape(word)}\b", text, flags=re.IGNORECASE)]
    transitions = extract_transition_phrases(text)
    passive_hits = 0
    tokenized = text.split()
    for index, token in enumerate(tokenized[:-1]):
        if token.lower().strip(",.;:()") in PASSIVE_MARKERS:
            next_token = tokenized[index + 1].lower().strip(",.;:()")
            if next_token.endswith("ed") or next_token in {"known", "given", "made", "taken"}:
                passive_hits += 1

    style_tags: list[str] = []
    if avg_sentence_length >= 28:
        style_tags.append("long-form sentences")
    elif avg_sentence_length >= 18:
        style_tags.append("measured sentence length")
    else:
        style_tags.append("compact sentence length")
    if transitions:
        style_tags.append("explicit connective logic")
    if date_count >= 2:
        style_tags.append("chronological framing")
    if citation_count:
        style_tags.append("citation-driven support")
    if passive_hits:
        style_tags.append("formal passive constructions")
    if rhetorical_move in {"factual_allegation", "numbered_allegation", "chronology"}:
        style_tags.append("fact-forward narration")

    return {
        "word_count": len(words),
        "sentence_count": len(sentences),
        "avg_sentence_length": avg_sentence_length,
        "citation_count": citation_count,
        "date_count": date_count,
        "transition_phrases": transitions,
        "qualifiers": qualifiers,
        "passive_voice_markers": passive_hits,
        "style_tags": style_tags,
    }


def style_summary(features: dict, rhetorical_move: str) -> str:
    elements = [rhetorical_move.replace("_", " ")]
    elements.extend(features["style_tags"][:4])
    if features["qualifiers"]:
        elements.append("qualifier-heavy precision")
    return ", ".join(dict.fromkeys(elements))


def infer_entity_label(entity: str) -> str:
    if entity.upper() == entity and len(entity.split()) >= 2:
        return "ENTITY"
    return "ENTITY"


def replace_with_placeholders(text: str, pattern: str, label: str) -> str:
    counter = 0
    mapping: dict[str, str] = {}

    def repl(match: re.Match[str]) -> str:
        nonlocal counter
        original = match.group(0)
        if original in mapping:
            return mapping[original]
        counter += 1
        placeholder = f"[[{label}_{counter}]]"
        mapping[original] = placeholder
        return placeholder

    return re.sub(pattern, repl, text)


def abstract_legal_text(text: str) -> str:
    text = normalize_paragraph_text(text)
    text = re.sub(
        r"\b(?:The\s+)?Defendant,\s+[A-Z][A-Za-z. ]+[A-Z],",
        "the Defendant,",
        text,
    )
    text = replace_with_placeholders(
        text,
        rf"\b(?:{MONTH_PATTERN})\s+\d{{1,2}},\s+\d{{4}}\b",
        "DATE",
    )
    text = replace_with_placeholders(text, r"\b\d{1,2}/\d{1,2}/\d{2,4}\b", "DATE")
    text = replace_with_placeholders(text, r"\b\d{1,2}:\d{2}\s*(?:a\.m\.|p\.m\.|AM|PM)?\b", "TIME")
    text = replace_with_placeholders(text, r"\b(?:Case|No\.)\s+[A-Za-z0-9:.-]+\b", "CASE_REF")
    text = replace_with_placeholders(text, r"\b\d+\s+U\.S\.C\.\s+§{1,2}\s*[\w().-]+\b", "STATUTE")
    text = replace_with_placeholders(text, r"\bCount\s+[A-Za-z0-9]+\b", "COUNT")

    title_patterns = [
        (r"\bUnited States of America\b", "the Government"),
        (r"\bUnited States\b", "the Government"),
        (r"\bThe Defendant\b", "the Defendant"),
        (r"\bDefendant\b", "the Defendant"),
    ]
    for pattern, replacement in title_patterns:
        text = re.sub(pattern, replacement, text)

    entity_pattern = re.compile(
        r"\b(?:[A-Z][a-z]+(?:\s+[A-Z]\.)?(?:\s+[A-Z][a-z]+){1,3}|[A-Z]{2,}(?:\s+[A-Z]{2,}){1,3})\b"
    )
    entity_map: dict[str, str] = {}
    entity_counter = 0

    def entity_repl(match: re.Match[str]) -> str:
        nonlocal entity_counter
        original = match.group(0)
        if original in LEGAL_ENTITY_EXCLUSIONS:
            return original
        if original.startswith("[[") and original.endswith("]]"):
            return original
        if re.fullmatch(MONTH_PATTERN, original, flags=re.IGNORECASE):
            return original
        if original in entity_map:
            return entity_map[original]
        entity_counter += 1
        placeholder = f"[[{infer_entity_label(original)}_{entity_counter}]]"
        entity_map[original] = placeholder
        return placeholder

    text = entity_pattern.sub(entity_repl, text)
    text = re.sub(r"\b\d{4,}\b", "[[NUMBER]]", text)
    text = re.sub(r"\bthe\s+the\s+Defendant\b", "the Defendant", text)
    text = re.sub(r"\bthe\s+the\s+Government\b", "the Government", text)
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\[\[ENTITY_\d+\]\],\s+the Defendant", "the Defendant", text)
    return text.strip()


def sentence_to_note(sentence: str) -> str:
    sentence = re.sub(r"^[A-Z][a-z]+,\s+", "", sentence)
    sentence = re.sub(r"\([^)]*\)", "", sentence)
    sentence = sentence.replace(";", ",")
    sentence = re.split(r",\s+(?:which|who|that|because|after|when)\b", sentence, maxsplit=1)[0]
    sentence = normalize_paragraph_text(sentence)
    words = sentence.split()
    if len(words) > 22:
        sentence = " ".join(words[:22]).rstrip(",.;:") + " ..."
    return sentence


def build_content_notes(text: str) -> list[str]:
    notes = []
    for sentence in split_sentences(text)[:4]:
        note = sentence_to_note(sentence)
        if re.fullmatch(r"\d+[.)]?", note):
            continue
        if note and note not in notes:
            notes.append(note)
    return notes


def build_training_instruction(author: str, rhetorical_move: str, style_summary_text: str) -> str:
    author_label = author or "the source author"
    return (
        f"Draft a {rhetorical_move.replace('_', ' ')} paragraph in the legal voice of {author_label}. "
        f"Preserve the source's disciplined structure, tone, and connective logic. "
        f"Target style: {style_summary_text}."
    )


def build_training_input(unit: ParagraphUnit, notes: list[str], features: dict) -> str:
    lines = [
        f"author: {unit.author}",
        f"document_type: {unit.document_type}",
        f"section_path: {' > '.join(unit.section_path) if unit.section_path else 'body'}",
        f"word_count_target: {features['word_count']}",
        f"sentence_count_target: {features['sentence_count']}",
        "style_cues: " + ", ".join(features["style_tags"][:4]),
        "drafting_notes:",
    ]
    lines.extend(f"- {note}" for note in notes)
    return "\n".join(lines)


def keep_style_record(unit: ParagraphUnit, rhetorical_move: str, features: dict) -> bool:
    if rhetorical_move == "caption":
        return False
    if features["word_count"] < 5:
        return False
    return True


def is_training_candidate(unit: ParagraphUnit, features: dict, min_words: int) -> bool:
    if features["word_count"] < min_words:
        return False
    if not unit.text or unit.text.isupper():
        return False
    if any(token in unit.text.lower() for token in ["united states district court", "defendant.", "grand jury original"]):
        return False
    return True


def build_style_record(unit: ParagraphUnit, document_id: str, features: dict,
                       rhetorical_move: str, summary: str, abstracted_text: str,
                       meta: dict) -> dict:
    return {
        "record_type": "style_analysis",
        "document_id": document_id,
        "source_file": unit.source_file,
        "paragraph_id": unit.paragraph_id,
        "author": unit.author,
        "document_type": unit.document_type,
        "page_start": unit.page_start,
        "ordinal": unit.ordinal,
        "section_path": unit.section_path,
        "rhetorical_move": rhetorical_move,
        "style_summary": summary,
        "style_features": features,
        "source_text": unit.text,
        "abstracted_text": abstracted_text,
        "ocr_used": meta.get("ocr_used", False),
    }


def build_training_pair(unit: ParagraphUnit, document_id: str, instruction: str,
                        input_text: str, output_text: str, rhetorical_move: str) -> dict:
    return {
        "instruction": instruction,
        "input": input_text,
        "output": output_text,
        "metadata": {
            "document_id": document_id,
            "source_file": unit.source_file,
            "paragraph_id": unit.paragraph_id,
            "author": unit.author,
            "document_type": unit.document_type,
            "page_start": unit.page_start,
            "section_path": unit.section_path,
            "rhetorical_move": rhetorical_move,
        },
    }


def write_jsonl(path: Path, records: Iterable[dict]) -> int:
    count = 0
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as handle:
        for record in records:
            handle.write(json.dumps(record, ensure_ascii=False) + "\n")
            count += 1
    return count


def process_document(path: Path, author_override: str | None, document_type_override: str | None,
                     ocr_mode: str, ocr_min_chars: int, ocr_lang: str,
                     min_words: int) -> tuple[list[dict], list[dict], dict]:
    preview_text = load_preview_text(path)
    author = author_override or "Unknown Author"
    document_type = detect_document_type(path, document_type_override, preview_text)

    if path.suffix.lower() == ".pdf":
        units, meta = parse_pdf_document(path, author, document_type, ocr_mode, ocr_min_chars, ocr_lang)
    elif path.suffix.lower() == ".docx":
        units, meta = parse_docx_document(path, author, document_type)
    else:
        units, meta = parse_text_document(path, author, document_type)

    style_records: list[dict] = []
    training_pairs: list[dict] = []
    move_counts: Counter[str] = Counter()
    document_id = meta["document_id"]

    for unit in units:
        rhetorical_move = classify_rhetorical_move(unit, unit.text)
        features = compute_style_features(unit.text, rhetorical_move)
        abstracted_text = abstract_legal_text(unit.text)
        summary = style_summary(features, rhetorical_move)

        if keep_style_record(unit, rhetorical_move, features):
            style_records.append(
                build_style_record(unit, document_id, features, rhetorical_move, summary, abstracted_text, meta)
            )
        move_counts[rhetorical_move] += 1

        if not is_training_candidate(unit, features, min_words):
            continue

        notes = build_content_notes(abstracted_text)
        if not notes:
            continue

        instruction = build_training_instruction(author, rhetorical_move, summary)
        input_text = build_training_input(unit, notes, features)
        training_pairs.append(
            build_training_pair(unit, document_id, instruction, input_text, abstracted_text, rhetorical_move)
        )

    stats = {
        "source_file": path.name,
        "document_id": document_id,
        "paragraphs": len(units),
        "style_records": len(style_records),
        "training_pairs": len(training_pairs),
        "ocr_used": meta.get("ocr_used", False),
        "ocr_pages": meta.get("ocr_pages", []),
        "rhetorical_moves": dict(move_counts),
    }
    return style_records, training_pairs, stats


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate legal style datasets from documents")
    parser.add_argument("input_path", help="A legal document file or directory")
    parser.add_argument("output_prefix", help="Output prefix for .style.jsonl and .pairs.jsonl")
    parser.add_argument("--author", help="Author label for the dataset")
    parser.add_argument("--document-type", help="Document type override, e.g. indictment")
    parser.add_argument("--min-words", type=int, default=50, help="Minimum words for a training pair")
    parser.add_argument(
        "--ocr-mode",
        choices=["auto", "always", "never"],
        default="auto",
        help="PDF OCR mode: auto uses OCR on pages with weak extracted text",
    )
    parser.add_argument(
        "--ocr-min-chars",
        type=int,
        default=250,
        help="In auto mode, OCR PDF pages whose extracted text is shorter than this many characters",
    )
    parser.add_argument("--ocr-lang", default="eng", help="Tesseract OCR language")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    input_path = Path(args.input_path)
    output_prefix = Path(args.output_prefix)
    files = collect_input_files(input_path)

    all_style_records: list[dict] = []
    all_training_pairs: list[dict] = []
    all_stats: list[dict] = []

    print(f"Scanning {input_path}")
    print(f"Found {len(files)} supported document(s)")

    for path in files:
        print(f"\nProcessing {path.name}...")
        style_records, training_pairs, stats = process_document(
            path=path,
            author_override=args.author,
            document_type_override=args.document_type,
            ocr_mode=args.ocr_mode,
            ocr_min_chars=args.ocr_min_chars,
            ocr_lang=args.ocr_lang,
            min_words=args.min_words,
        )
        all_style_records.extend(style_records)
        all_training_pairs.extend(training_pairs)
        all_stats.append(stats)

        ocr_note = f" OCR pages: {stats['ocr_pages']}" if stats["ocr_pages"] else ""
        print(
            f"  paragraphs={stats['paragraphs']} style_records={stats['style_records']} "
            f"training_pairs={stats['training_pairs']} ocr_used={stats['ocr_used']}" + ocr_note
        )

    style_path = output_prefix.with_suffix(".style.jsonl")
    pairs_path = output_prefix.with_suffix(".pairs.jsonl")
    summary_path = output_prefix.with_suffix(".summary.json")

    style_count = write_jsonl(style_path, all_style_records)
    pair_count = write_jsonl(pairs_path, all_training_pairs)

    summary = {
        "input_path": str(input_path),
        "output_prefix": str(output_prefix),
        "style_records": style_count,
        "training_pairs": pair_count,
        "documents": all_stats,
    }
    summary_path.parent.mkdir(parents=True, exist_ok=True)
    summary_path.write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")

    print(f"\nOutput style dataset: {style_path}")
    print(f"Output training pairs: {pairs_path}")
    print(f"Output summary: {summary_path}")
    print(f"Total style records: {style_count}")
    print(f"Total training pairs: {pair_count}")


if __name__ == "__main__":
    main()